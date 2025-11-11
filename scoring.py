# -*- coding: utf-8 -*-
"""
Created on Tue Nov  4 13:42:47 2025
@author: Chiara Aquino
"""

import os
import numpy as np
import pandas as pd
from utils.file_ops import sanitize_filename
from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill, Alignment
import re
from tkinter import messagebox

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def safe_save_docx(save_func, *args, **kwargs):
    """
    Safely save Word (.docx) files and show a friendly message if the file is already open.
    Example:
        safe_save_docx(document.save, "path.docx")
    """
    try:
        return save_func(*args, **kwargs)
    except PermissionError:
        messagebox.showerror(
            "File in Use",
            "❌ Cannot save the Word report.\n\n"
            "Please close the file before trying again."
        )
        raise



def safe_save_excel(save_func, *args, **kwargs):
    """
    Safely save Excel files and show a friendly message if the file is already open.
    Example:
        safe_save_excel(wb.save, "path.xlsx")
        safe_save_excel(df.to_excel, "path.xlsx", index=False)
    """
    try:
        return save_func(*args, **kwargs)
    except PermissionError:
        messagebox.showerror(
            "File in Use",
            "❌ Cannot save one or more files.\n\n"
            "Please close all Excel files before trying again."
        )
        raise



def balanced_assignments(application_ids, scorers, reviews_per_app):
    """
    Fair round-robin assignment.
    Each application is reviewed `reviews_per_app` times.
    Workload difference across scorers <= 1.
    """
    import random
    random.shuffle(application_ids)  # keep random element so each run is a bit different

    assignments = {s: [] for s in scorers}
    app_assignments = {a: [] for a in application_ids}

    # Keep assigning each application to the least-loaded scorers
    for app in application_ids:
        for _ in range(reviews_per_app):
            # sort scorers by how many apps they've already been assigned
            sorted_scorers = sorted(scorers, key=lambda s: len(assignments[s]))
            for scorer in sorted_scorers:
                # ensure each app gets unique scorers
                if scorer not in app_assignments[app]:
                    assignments[scorer].append(app)
                    app_assignments[app].append(scorer)
                    break

    # ✅ sanity check
    loads = {s: len(a) for s, a in assignments.items()}
    max_load = max(loads.values())
    min_load = min(loads.values())
    print(f"⚖️ Load distribution (balanced): {loads}")
    print(f"Difference between max/min: {max_load - min_load}")

    return assignments, app_assignments




def make_matrix_and_scorecards(applications_file, scorers, output_folder, reviews_per_app, scorecard_columns):
    """
    Build the assignment matrix and scorer Excel files.
    """
    # === Load applications ===
    applications_df = pd.read_excel(applications_file)
    
    if applications_df.shape[1] < 2:
        raise ValueError("The eligible applications file must have at least two columns (ID and group name).")
    
    # Identify first two columns as ID + Name
    id_col = applications_df.columns[0]
    group_col = applications_df.columns[1]
    
    # If there’s a “Topic” column, keep it but ignore it for simple scoring
    extra_cols = [c for c in applications_df.columns[2:] if c.lower() == "topic"]
    if extra_cols:
        print(f"⚠️ Detected extra columns in file (ignored for scoring): {extra_cols}")
    
    # --- Clean ID column safely ---
    id_series = applications_df[id_col].astype(str).fillna("")
    applications_df[id_col] = id_series.apply(
        lambda x: x.zfill(3) if x.isdigit() else x
    )
    
    # --- Continue with only first two columns for assignments ---
    applications_core = applications_df[[id_col, group_col]].copy()
    application_ids = applications_core[id_col].dropna().unique().tolist()
    
    if not application_ids:
        raise ValueError("No valid application IDs found in the eligible applications file.")


    application_ids = applications_df[id_col].unique().tolist()
    scorers = [s.strip() for s in scorers if str(s).strip()]

    if not scorers:
        raise ValueError("Please add at least one scorer.")
    if len(application_ids) == 0:
        raise ValueError("No applications found in the eligible applications file.")
    if reviews_per_app not in (1, 2, 3):
        raise ValueError("Reviews per application must be 1, 2 or 3.")

    # === Fair assignment ===
    assignments, app_assignments = balanced_assignments(application_ids, scorers, reviews_per_app)

    # === Matrix summary ===
    matrix_data = [{"Name": s, "Total applications per scorer": len(a)} for s, a in assignments.items()]
    matrix_df = pd.DataFrame(matrix_data)
    matrix_df.loc[len(matrix_df)] = {
        "Name": "Total assigned",
        "Total applications per scorer": matrix_df["Total applications per scorer"].sum()
    }

    os.makedirs(output_folder, exist_ok=True)
    matrix_path = os.path.join(output_folder, "Matrix.xlsx")
    try:
        safe_save_excel(matrix_df.to_excel, matrix_path, index=False)
    except PermissionError:
        return

    
    # === After saving Matrix.xlsx ===
    report_path = generate_summary_report(
        matrix_df, applications_df, scorers, reviews_per_app, output_folder
    )
    print(f"✅ Summary report created")


    # === Create individual scorecards ===
    scorers_folder = os.path.join(output_folder, "Scorers")
    os.makedirs(scorers_folder, exist_ok=True)

    if not scorecard_columns:
        scorecard_columns = ["Points out of 6", "Comments"]

    error_occurred = False  # track if any save failed

    for scorer, apps in assignments.items():
        if not apps:
            continue
    
        df = applications_df[applications_df[id_col].isin(apps)][[id_col, group_col]].copy()
        for col in scorecard_columns:
            if col not in df.columns:
                df[col] = ""
    
        ordered_cols = [id_col, group_col] + [c for c in scorecard_columns]
        df = df[ordered_cols]
    
        scorer_folder = os.path.join(scorers_folder, scorer)
        os.makedirs(scorer_folder, exist_ok=True)
        file_path = os.path.join(scorer_folder, f"{scorer} Scorecard.xlsx")
    
        try:
            safe_save_excel(df.to_excel, file_path, index=False, engine="openpyxl")
        except PermissionError:
            messagebox.showerror(
                "File in Use",
                f"❌ Cannot save '{os.path.basename(file_path)}'.\n\n"
                "Please close it in Excel before trying again."
            )
            error_occurred = True
            break  # stop assigning further scorers


        # === Formatting ===
        wb = load_workbook(file_path)
        ws = wb.active
        ws.freeze_panes = "C2"

        header_font = Font(name="Arial", bold=True, size=11)
        row_font = Font(name="Arial", size=11)
        group_font = Font(name="Arial", size=13)
        fill = PatternFill(start_color="FFFFCC", end_color="FFFFCC", fill_type="solid")
        wrap = Alignment(wrap_text=True, horizontal="center", vertical="top")
        center = Alignment(horizontal="center", vertical="center")

        for cell in ws[1]:
            cell.font = header_font
            cell.fill = fill
            cell.alignment = wrap
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, max_col=ws.max_column):
            for cell in row:
                cell.font = row_font
                cell.alignment = center
        for cell in ws["A"]:
            if cell.row > 1:
                cell.font = group_font

        for col in ws.columns:
            maxlen = max((len(str(c.value)) if c.value else 0) for c in col)
            ws.column_dimensions[col[0].column_letter].width = maxlen + 2

        try:
            safe_save_excel(wb.save, file_path)
        except PermissionError:
            messagebox.showerror(
                "File in Use",
                f"❌ Cannot save formatted file '{os.path.basename(file_path)}'.\n\n"
                "Please close it before trying again."
            )
            error_occurred = True
            break
        

    # === Summary printout ===
    total_apps = len(application_ids)
    total_scorers = len(assignments)
    folders_created = sum(1 for s, apps in assignments.items() if apps)
    
def generate_summary_report(matrix_df, applications_df, scorers, reviews_per_app, output_folder):
    """
    Create a Word report summarizing the scoring setup.
    """
    report_path = os.path.join(output_folder, "Scorecard Summary Report.docx")
    document = Document()

    # Title
    title = document.add_heading("Edge Fund Scoring Summary Report", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("")
    document.add_paragraph("📅 Report generated automatically by the Edgy Scoring Module").italic = True

    total_apps = len(applications_df)
    total_scorers = len(scorers)
    total_reviews = total_apps * reviews_per_app
    # Compute actual average from matrix
    if "Total applications per scorer" in matrix_df.columns:
        avg_per_scorer = round(matrix_df["Total applications per scorer"][:-1].mean(), 2)
    else:
        avg_per_scorer = round(total_reviews / total_scorers, 2)

    # Summary section
    document.add_paragraph(f"• Total eligible applications: {total_apps}")
    document.add_paragraph(f"• Total scorers: {total_scorers}")
    document.add_paragraph(f"• Reviews per application: {reviews_per_app}")
    expected_reviews = total_apps * reviews_per_app
    document.add_paragraph(f"• Expected number of applications to review = {total_apps} x {reviews_per_app} = {expected_reviews}")

    document.add_paragraph(f"• Average workload per scorer: {avg_per_scorer}")

    document.add_paragraph("\n")

    # Add a nice section box title
    p = document.add_paragraph()
    run = p.add_run("Assignment Matrix Summary")
    run.bold = True
    run.font.size = Pt(13)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Create matrix table
    table = document.add_table(rows=1, cols=len(matrix_df.columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for j, col_name in enumerate(matrix_df.columns):
        hdr_cells[j].text = str(col_name)
        hdr_cells[j].paragraphs[0].runs[0].bold = True

    for _, row in matrix_df.iterrows():
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = str(val)

    # Auto fit table width
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(2)


    try:
        safe_save_docx(document.save, report_path)
    except PermissionError:
        print("❌ Report could not be saved because it is open in Word.")
        return

    print(f"📄 Report saved at {report_path}")
    return report_path


def make_topic_based_scorecards(
    eligible_file,
    output_folder,
    scorers_topics,
    id_column,
    topic_column,
    group_col_eligible,
    general_topics,
    reviews_per_app,
    scorecard_columns,
    create_report=True,
    progress_callback=None
):
    """
    Build topic-based scorecards directly from the Eligible Applications file.
    - Includes ID + Group Name columns (no Topics)
    - Assigns apps only based on topic match (not just shared topics)
    - Prioritizes scorers who chose General topics, then rebalances others if overloaded
    - Unassigned apps remain unassigned (reported)
    - Balanced workload across scorers
    - Creates 'Applications by Topics.xlsx' with one sheet per topic
    - Tracks exact topic reason for each assignment
    """

    # === Load data ===
    eligible_df = pd.read_excel(eligible_file)
    for col in [id_column, group_col_eligible, topic_column]:
        if col not in eligible_df.columns:
            raise ValueError(f"Column '{col}' not found in eligible file.")

    eligible_df[group_col_eligible] = eligible_df[group_col_eligible].astype(str).apply(sanitize_filename)

    # --- Ensure unique and clean group names ---
    duplicates_warnings = []
    eligible_df[group_col_eligible] = eligible_df[group_col_eligible].astype(str).str.strip()
    if eligible_df[group_col_eligible].duplicated().any():
        msg = "⚠️ Warning: Duplicate group names found! Please check your eligible file."
        print(msg)
        duplicates_warnings.append(msg)
        duplicates = eligible_df[eligible_df[group_col_eligible].duplicated(keep=False)]
        dup_str = duplicates[[id_column, group_col_eligible, topic_column]].to_string(index=False)
        duplicates_warnings.append(dup_str)


    # --- Parse topics ---
    def parse_topics(x):
        if pd.isna(x):
            return []
        return [t.strip().lower() for t in re.split(r",\s*", str(x)) if t.strip()]

    eligible_df["parsed_topics"] = eligible_df[topic_column].apply(parse_topics)

    # === Prepare structures ===
    assignments = {s: [] for s in scorers_topics}
    assignment_reasons = {s: {} for s in scorers_topics}  # NEW: topic reason for assignment
    app_assignments = {row[group_col_eligible]: [] for _, row in eligible_df.iterrows()}

    total_apps = len(eligible_df)
    def get_load(scorer): return len(assignments.get(scorer, []))

    # === Assignment loop ===
    for idx, (_, row) in enumerate(eligible_df.iterrows(), start=1):
        app_name = row[group_col_eligible]
        app_topics = {t.lower().strip() for t in row["parsed_topics"]}
        is_general_app = any(gt.lower().strip() in app_topics for gt in general_topics)

        eligible_pairs = []  # (scorer, reason_topic)

        # CASE 1: Normal (no general topic)
        if not is_general_app:
            for scorer, st_topics in scorers_topics.items():
                st_lower = {t.lower().strip() for t in st_topics}
                matched = list(st_lower & app_topics)
                if matched:
                    eligible_pairs.append((scorer, matched[0]))

        # CASE 2: General app
        # --- CASE 2: App has a general topic (improved priority logic)
        else:
            matched_topic_scorers = []
            general_scorers = []
            balancing_scorers = []
        
            for scorer, st_topics in scorers_topics.items():
                st_lower = {t.lower().strip() for t in st_topics}
        
                # ✅ 1️⃣ Prefer scorers who match non-general topics in this app
                matched_non_general = [t for t in app_topics if t not in general_topics and t in st_lower]
                if matched_non_general:
                    matched_topic_scorers.append((scorer, matched_non_general[0]))
                    continue
        
                # ✅ 2️⃣ Then scorers who explicitly chose general topics
                if any(gt.lower().strip() in st_lower for gt in general_topics):
                    general_scorers.append((scorer, "general"))
                    continue
        
                # ✅ 3️⃣ Lastly, low-load scorers (fairness fallback)
                balancing_scorers.append((scorer, "general_balance"))
        
            # Sort by current load to balance work
            matched_topic_scorers.sort(key=lambda x: get_load(x[0]))
            general_scorers.sort(key=lambda x: get_load(x[0]))
            balancing_scorers.sort(key=lambda x: get_load(x[0]))
            
            eligible_pairs = matched_topic_scorers + general_scorers + balancing_scorers
        
        # --- If no scorers qualify, leave unassigned
        if not eligible_pairs:
            continue

        # --- Sort by load and assign ---
        eligible_pairs.sort(key=lambda x: get_load(x[0]))
        
        assigned_count = 0
        assignment_warnings = []  # <-- collect diagnostics to add to report
        
        for s, reason_topic in eligible_pairs:
            if s not in assignments:
                msg = f"⚠️ Scorer '{s}' not found in assignments dictionary. Skipped for app '{app_name}'."
                print(msg)
                assignment_warnings.append(msg)
                continue
        
            # Ensure app not already assigned to this scorer
            if app_name in assignments[s]:
                continue
        
            # Perform the assignment
            assignments[s].append(app_name)
            app_assignments[app_name].append(s)
            assignment_reasons[s][app_name] = reason_topic
            assigned_count += 1
        
            if assigned_count >= reviews_per_app:
                break
        
        # --- Post-check: not enough scorers found
        if assigned_count < reviews_per_app:
            msg = f"❌ App '{app_name}' assigned only {assigned_count}/{reviews_per_app} reviewers. Only {len(eligible_pairs)} eligible scorers found."
            print(msg)
            assignment_warnings.append(msg)


        if progress_callback:
            progress_callback(idx, total_apps)

    # === Build Matrix by Topic (count only topic reason per assignment) ===
    os.makedirs(output_folder, exist_ok=True)
    topic_list = sorted({t for lst in eligible_df["parsed_topics"] for t in lst})
    matrix_data = []
    for scorer, apps in assignments.items():
        row = {"Scorer": scorer}
    
        for t in topic_list:
            count = 0
            for app in assignments[scorer]:
                reason = assignment_reasons[scorer].get(app)
    
                # 1️⃣ Normal topic match
                if reason == t:
                    count += 1
    
                # 2️⃣ General topic balancing: count under actual general topic
                elif reason in ("general", "general_balance") and any(
                    gt.lower().strip() == t for gt in general_topics
                ):
                    count += 1
    
            row[t] = count

        # Total applications assigned to this scorer
        row["Total apps"] = len(apps)
        matrix_data.append(row)

    
    matrix_df = pd.DataFrame(matrix_data)
    matrix_df = matrix_df[["Scorer"] + topic_list + ["Total apps"]]
    totals = {"Scorer": "Total"}
    for t in topic_list:
        totals[t] = matrix_df[t].sum()
    totals["Total apps"] = matrix_df["Total apps"].sum()
    matrix_df.loc[len(matrix_df)] = totals
    matrix_path = os.path.join(output_folder, "Matrix by Topic.xlsx")
    try:
        safe_save_excel(matrix_df.to_excel, matrix_path, index=False)
    except PermissionError:
        return
    print(f"📊 Matrix saved at: {matrix_path}")

        # === Applications by Topics workbook (safe sheet names) ===
    topics_path = os.path.join(output_folder, "Applications by Topics.xlsx")

    def sanitize_sheet_name(name: str) -> str:
        invalid_chars = r'[:\\/?*\[\]]'
        safe = re.sub(invalid_chars, '', name)
        safe = safe.strip() or "Untitled"
        return safe[:31]

    try:
        with pd.ExcelWriter(topics_path, engine="openpyxl") as writer:
            for topic in topic_list:
                topic_apps = eligible_df[
                    eligible_df["parsed_topics"].apply(lambda lst: topic in lst)
                ][[id_column, group_col_eligible, topic_column]]
                topic_apps.to_excel(writer, sheet_name=sanitize_sheet_name(topic), index=False)
        print(f"📘 Applications by Topics saved at: {topics_path}")

    except PermissionError:
        messagebox.showerror(
            "File in Use",
            f"❌ Cannot save '{os.path.basename(topics_path)}'.\n\n"
            "Please close it in Excel before trying again."
        )
        return

    # === ⚡ NEW: Create 'Assignments by Application (Topics).xlsx' ===
    try:
        max_scorers = max((len(v) for v in app_assignments.values()), default=0)
        rows = []

        for _, row in eligible_df.iterrows():
            app_id = row[id_column]
            app_name = row[group_col_eligible]
            topics_str = row[topic_column]
            scorers_for_app = app_assignments.get(app_name, [])
            record = {
                id_column: app_id,
                group_col_eligible: app_name,
                topic_column: topics_str
            }
            # Add Scorer 1, Scorer 2, ... dynamically
            for i in range(max_scorers):
                col_name = f"Scorer {i+1}"
                record[col_name] = scorers_for_app[i] if i < len(scorers_for_app) else ""
            rows.append(record)

        assignments_df = pd.DataFrame(rows)
        assign_path = os.path.join(output_folder, "Scorer Assignments.xlsx")
        safe_save_excel(assignments_df.to_excel, assign_path, index=False)
        print(f"📒 Scorer Assignment saved at: {assign_path}")

    except Exception as e:
        print(f"⚠️ Could not create 'Scorer Assignments).xlsx': {e}")


    except PermissionError:
        messagebox.showerror(
            "File in Use",
            f"❌ Cannot save '{os.path.basename(topics_path)}'.\n\n"
            "Please close it in Excel before trying again."
        )
        return


    # === Identify diagnostics ===
    assigned_apps = {a for lst in assignments.values() for a in lst}
    left_out_apps = [a for a in eligible_df[group_col_eligible] if a not in assigned_apps]
    app_review_counts = {a: len(r) for a, r in app_assignments.items()}
    underassigned = {a: c for a, c in app_review_counts.items() if c < reviews_per_app and c > 0}

    # === Off-topic check (should be none)
    off_topic_assignments = []
    app_topics_map = {row[group_col_eligible]: row["parsed_topics"] for _, row in eligible_df.iterrows()}
    for scorer, apps in assignments.items():
        chosen_topics = {t.lower().strip() for t in scorers_topics.get(scorer, [])}
        for app in apps:
            app_topics = set(app_topics_map.get(app, []))
            reason = assignment_reasons[scorer].get(app)
            if reason not in app_topics and reason not in general_topics:
                off_topic_assignments.append((scorer, app, reason))

    # === Create individual scorecards ===
    scorers_folder = os.path.join(output_folder, "Scorers")
    os.makedirs(scorers_folder, exist_ok=True)
    error_occurred = False  # track if any save failed

    for scorer, apps in assignments.items():
        if not apps:
            continue
    
        subset = eligible_df[eligible_df[group_col_eligible].isin(apps)][[id_column, group_col_eligible]].copy()
        subset.rename(columns={id_column: "ID", group_col_eligible: "Group name"}, inplace=True)
        for c in scorecard_columns:
            subset[c] = ""
    
        scorer_folder = os.path.join(scorers_folder, sanitize_filename(scorer))
        os.makedirs(scorer_folder, exist_ok=True)
        out_path = os.path.join(scorer_folder, f"{sanitize_filename(scorer)} Scorecard.xlsx")
    
        try:
            safe_save_excel(subset.to_excel, out_path, index=False, engine="openpyxl")
        except PermissionError:
            messagebox.showerror(
                "File in Use",
                f"❌ Cannot save '{os.path.basename(out_path)}'.\n\n"
                "Please close it in Excel before trying again."
            )
            error_occurred = True
            break

        # === Formatting (reuse same style as make_matrix_and_scorecards) ===
        wb = load_workbook(out_path)
        ws = wb.active
        ws.freeze_panes = "C2"
        
        header_font = Font(name="Arial", bold=True, size=11)
        row_font = Font(name="Arial", size=11)
        group_font = Font(name="Arial", size=13)
        fill = PatternFill(start_color="FFFFCC", end_color="FFFFCC", fill_type="solid")
        wrap = Alignment(wrap_text=True, horizontal="center", vertical="top")
        center = Alignment(horizontal="center", vertical="center")
        
        for cell in ws[1]:
            cell.font = header_font
            cell.fill = fill
            cell.alignment = wrap
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row, max_col=ws.max_column):
            for cell in row:
                cell.font = row_font
                cell.alignment = center
        for cell in ws["A"]:
            if cell.row > 1:
                cell.font = group_font
        
        for col in ws.columns:
            maxlen = max((len(str(c.value)) if c.value else 0) for c in col)
            ws.column_dimensions[col[0].column_letter].width = maxlen + 2
        
        try:
            safe_save_excel(wb.save, out_path)
        except PermissionError:
            messagebox.showerror(
                "File in Use",
                f"❌ Cannot save formatted file '{os.path.basename(out_path)}'.\n\n"
                "Please close it before trying again."
            )
            error_occurred = True
            break



    # === Unified Report ===
    if create_report:
        from scoring import generate_summary_report
        matrix_summary_data = [{"Name": s, "Total applications per scorer": len(a)} for s, a in assignments.items()]
        matrix_summary = pd.DataFrame(matrix_summary_data)
        matrix_summary.loc[len(matrix_summary)] = {
            "Name": "Total assigned",
            "Total applications per scorer": matrix_summary["Total applications per scorer"].sum()
        }
    
        report_path = generate_summary_report(
            matrix_summary,
            eligible_df,
            list(scorers_topics.keys()),
            reviews_per_app,
            output_folder
        )
        doc = Document(report_path)
    
        # === Applications not or under marked ===
        doc.add_heading("Applications Not or Under Marked", level=1)

        if left_out_apps:
            doc.add_paragraph(f"⚠️ {len(left_out_apps)} applications were not assigned to any scorer:", style="List Bullet")
            for app in left_out_apps:
                doc.add_paragraph(str(app), style="List Bullet")
        else:
            doc.add_paragraph("✅ All applications were assigned at least once.")
    
        if underassigned:
            doc.add_paragraph(f"⚠️ {len(underassigned)} applications were assigned fewer than {reviews_per_app} reviews:", style="List Bullet")
            for app_id, count in underassigned.items():
                doc.add_paragraph(f"{app_id}: {count} reviews assigned", style="List Bullet")
        
        if assignment_warnings:
            doc.add_heading("Assignment Warnings", level=1)
            for msg in assignment_warnings:
                doc.add_paragraph(msg, style="List Bullet")
        if duplicates_warnings:
            doc.add_heading("Duplicate Applications", level=1)
            for msg in duplicates_warnings:
                doc.add_paragraph(msg, style="List Bullet")
                
        else:
            doc.add_paragraph("✅ All applications have the required number of reviews.")
    
        # === Cross-topic / Off-topic Assignments (improved detail) ===
        doc.add_heading("Cross-Topic (Off-Topic) Assignments", level=1)
        if off_topic_assignments:
            doc.add_paragraph(f"⚠️ {len(off_topic_assignments)} cross-topic assignments detected:")
            for scorer, app, reason in off_topic_assignments:
                scorer_prefs = ", ".join(scorers_topics.get(scorer, []))
                app_topics = ", ".join(app_topics_map.get(app, []))
                doc.add_paragraph(
                    f"{scorer} (preferred: {scorer_prefs}) → {app} (topics: {app_topics}) "
                    f"(reason: {reason})",
                    style="List Bullet"
                )
        else:
            doc.add_paragraph("✅ No scorers were assigned applications outside their chosen topics.")
    
        # === Save final report ===
        try:
            safe_save_docx(doc.save, report_path)
        except PermissionError:
            print("❌ Report could not be saved because it is open in Word.")
            return

        print(f"📄 Report saved")


    print("✅ Topic-based scorecards created successfully with fair and topic-specific assignment.")
