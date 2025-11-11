# -*- coding: utf-8 -*-
"""
Created on Sun Nov  2 15:03:18 2025

@author: Chiara Aquino
"""

import os, sys
import re
import pandas as pd
from docx import Document
from utils.file_ops import sanitize_filename
import numpy as np
from tkinter import messagebox
import json



def resource_path(relative_path):
    """Get absolute path to resource, works for dev and for PyInstaller."""
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


def get_external_logic_path():
    """
    Preferred external path next to the executable or script,
    so users can edit the file even if running as a frozen exe.
    """
    # If frozen (exe), sys.executable points to the exe; otherwise use __file__ dir
    if getattr(sys, "frozen", False):
        exe_dir = os.path.dirname(sys.executable)
    else:
        exe_dir = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(exe_dir, "logic_config.json")

def load_logic_config():
    """
    Loads logic configuration. Priority:
      1) external editable file next to exe/script (get_external_logic_path())
      2) bundled default via resource_path (useful if app bundled and default provided)
    Raises FileNotFoundError if neither exists.
    """
    # 1) try external editable file next to exe/script
    ext_path = get_external_logic_path()
    if os.path.exists(ext_path):
        with open(ext_path, "r", encoding="utf-8") as f:
            return json.load(f)

    # 2) fallback to bundled default (resource_path) (read-only inside bundle)
    bundled_path = resource_path("logic_config.json")
    if os.path.exists(bundled_path):
        with open(bundled_path, "r", encoding="utf-8") as f:
            return json.load(f)

    raise FileNotFoundError(f"logic_config.json not found at either {ext_path} or {bundled_path}")

def save_logic_config(updated_logic):
    """
    Save to the external editable config file next to exe/script.
    Use this from your GUI editor to persist changes.
    """
    ext_path = get_external_logic_path()
    with open(ext_path, "w", encoding="utf-8") as f:
        json.dump(updated_logic, f, indent=2)

def lookup_result_for_answers(answers, logic):
    """
    Looks up the result for the given list of answers (e.g. ['yes','unsure','no'])
    using the provided logic dictionary (from logic_config.json).

    This version ignores order — it compares counts of yes/no/unsure
    to find a matching rule.
    """
    # Normalize answers
    answers = [a.strip().lower() for a in answers if a.strip()]

    # Case 1: no answers at all
    if not answers:
        return "yes"   # or "no" if you prefer stricter default

    # Case 2: one scorer only
    if len(answers) == 1:
        single = answers[0]
        # Special rule: if single scorer says "unsure", treat as "yes"
        if single == "unsure":
            return "yes"
        return single

    # Case 3: two scorers — use 2-scorer logic
    if len(answers) == 2:
        rule_dict = logic.get("LOGIC_2SCORERS", {})
    else:
        # Three or more scorers → use 3-scorer logic
        rule_dict = logic.get("LOGIC_3SCORERS", {})

    # Compute counts for this set of answers
    counts = {
        "yes": answers.count("yes"),
        "no": answers.count("no"),
        "unsure": answers.count("unsure"),
    }

    # Try to find a matching rule (order-insensitive)
    for key, result in rule_dict.items():
        key_answers = [k.strip().lower() for k in key.split(",") if k.strip()]
        key_counts = {
            "yes": key_answers.count("yes"),
            "no": key_answers.count("no"),
            "unsure": key_answers.count("unsure"),
        }
        if key_counts == counts:
            return result

    # Default fallback if no rule matches
    return "yes"





def create_application_cards(responses_path, columns_to_select, output_folder, group_col,
                              progress_callback=None, num_apps=None, topics_col=None):
    """
    Creates Word documents (Admin View + Scorers View) from an Excel/CSV responses file.

    Parameters:
        responses_path (str): Path to the Excel/CSV file with responses
        columns_to_select (list): Columns selected for Scorers View
        output_folder (str): Folder where to save the results
        group_col (str): Column name selected by user as "organisation name"
        progress_callback (function): Optional progress update callback
    """

    # --- Read file ---
    df = pd.read_excel(responses_path) if responses_path.endswith(('.xlsx', '.xls')) else pd.read_csv(responses_path)
    
    # --- Limit number of applications if requested ---
    if num_apps is not None and isinstance(num_apps, int):
        df = df.head(num_apps)

    # --- Handle timestamp column safely ---
    if 'Timestamp' in df.columns:
        df['Timestamp'] = pd.to_datetime(df['Timestamp'], errors='coerce')
        df['Unique Application ID'] = df['Timestamp'].astype('int64') // 10**9
    else:
        # Fallback if no timestamp column
        df['Unique Application ID'] = range(1, len(df) + 1)

    # --- Prepare output folders ---
    scorers_path = os.path.join(output_folder, "Scorers View")
    admin_path = os.path.join(output_folder, "Admin View")
    os.makedirs(scorers_path, exist_ok=True)
    os.makedirs(admin_path, exist_ok=True)

    # --- Clean up columns_to_select to ensure no duplicates ---
    columns_to_select = list(dict.fromkeys(columns_to_select))  # preserve order but remove duplicates

    total = len(df)
    count = 0
    
    # collect ID + Group + optional Topics for later use in marking sheet
    topics_data = []

    # --- Iterate through applications ---
    for _, row in df.iterrows():
        raw_group_name = str(row.get(group_col, "Unknown Group")).strip()
        group_name = sanitize_filename(raw_group_name)

        # Use the same sequential ID used in filenames (001, 002, ...)
        uid_str = f"{count+1:03d}"

        # --- ADMIN VIEW ---
        doc_admin = Document()
        doc_admin.add_heading(raw_group_name, level=1)
        for col_name in df.columns:
            p = doc_admin.add_paragraph()
            p.add_run(f"{col_name}: ").bold = True
            p.add_run(str(row[col_name]))
        admin_filename = os.path.join(admin_path, f"{uid_str}_{group_name}.docx")
        doc_admin.save(admin_filename)

        # --- SCORERS VIEW ---
        if columns_to_select:
            df_scorer = row[columns_to_select]
        else:
            df_scorer = row

        doc_scorer = Document()
        doc_scorer.add_heading(raw_group_name, level=1)
        for col_name in df_scorer.index:
            p = doc_scorer.add_paragraph()
            p.add_run(f"{col_name}: ").bold = True
            p.add_run(str(df_scorer[col_name]))
        scorer_filename = os.path.join(scorers_path, f"{uid_str}_{group_name}.docx")
        doc_scorer.save(scorer_filename)

        # --- NEW: store mapping for eligibility sheet ---
        if topics_col and topics_col in df.columns:
            topics_value = row.get(topics_col, "")
            topics_data.append({
                "Unique Application ID": uid_str,
                "Group Name": raw_group_name,
                "Topics": topics_value
            })
        else:
            topics_data.append({
                "Unique Application ID": uid_str,
                "Group Name": raw_group_name
            })

        count += 1
        if progress_callback:
            progress_callback(count, total)

    # --- NEW: save helper Excel with IDs, names and optional Topics ---
    if topics_data:
        topics_df = pd.DataFrame(topics_data)
        summary_file = os.path.join(output_folder, "Application Summary.xlsx")
        topics_df.to_excel(summary_file, index=False)


def convert_to_word(df, output_path, admin=True):
    count = 0
    for _, row in df.iterrows():
        doc = Document()
        group_name = str(row.get('Unique Application ID', 'NA')) + "_" + str(row.get('Group Name', 'Unknown'))
        doc.add_heading(group_name, level=1)

        for col_name in df.columns:
            p = doc.add_paragraph()
            p.add_run(f"{col_name}: ").bold = True
            p.add_run(str(row[col_name]))

        filename = sanitize_filename(group_name) + ".docx"
        doc.save(os.path.join(output_path, filename))
        count += 1

    print(f"Total Word documents created ({'Admin' if admin else 'Scorers'} View): {count}")
    


def convert_zero(x):
    """Standardise cell values: keep 'yes'/'no'/'unsure' or 0."""
    if pd.isna(x) or str(x).strip() == "" or str(x).strip() == "0":
        return 0
    s = str(x).strip().lower()
    if s in ["yes", "no", "unsure"]:
        return s
    return 0



def mark_eligibility(indata, applications, summary_output):


    logic = load_logic_config()

    # --- Identify columns ---
    app_id_col = indata.columns[0]
    org_name_col = indata.columns[1]

    answer_col_patterns = re.compile(r'\byes\b.*\bno\b|\byes\s*/\s*no\b|\byes/ ?no\b', flags=re.I)
    all_cols = list(indata.columns)

    remaining_cols = all_cols[2:]
    # NEW: exclude Topics column(s) from the answer columns
    remaining_cols = [c for c in remaining_cols if "topic" not in str(c).lower()]

    answer_cols = [c for c in remaining_cols if answer_col_patterns.search(str(c))]
    if not answer_cols:
        answer_cols = remaining_cols

    # --- Normalize ---
    df = indata.copy()
    df[answer_cols] = df[answer_cols].fillna("").applymap(lambda x: str(x).strip().lower())
    df.set_index(app_id_col, inplace=True)

    df["Total Scorers"] = df[answer_cols].astype(bool).sum(axis=1)
    df["Number of Unsure"] = (df[answer_cols] == "unsure").sum(axis=1)
    df["Final Score"] = ""

    # --- Apply logic based on number of scorers ---
    for i, row in df.iterrows():
        vals = [row[c] for c in answer_cols if row[c]]
        df.at[i, "Final Score"] = lookup_result_for_answers(vals, logic)

    df_final = df.reset_index()

    # --- Merge back ---
    needed_merge_cols = [app_id_col, "Final Score", "Total Scorers", "Number of Unsure"]
    needed_merge_cols = [c for c in needed_merge_cols if c in df_final.columns]
    df_merged = pd.merge(indata, df_final[needed_merge_cols], on=app_id_col, how="left").fillna("")
    df_merged["Final Score"] = df_merged["Final Score"].astype(str).str.lower()

    # --- Split results ---
    eligible = df_merged[df_merged["Final Score"] == "yes"]
    noneligible = df_merged[df_merged["Final Score"] == "no"]

    summary = df_merged[[org_name_col] + [c for c in ["Final Score", "Total Scorers", "Number of Unsure"] if c in df_merged.columns]]

    # --- Save outputs ---
    base_dir = os.path.dirname(summary_output)
    summary_path = os.path.join(base_dir, "Eligibility Summary.xlsx")
    eligible_path = os.path.join(base_dir, "Eligible Applications.xlsx")
    noneligible_path = os.path.join(base_dir, "Noneligible Applications.xlsx")
    
    # Ensure the first column (Unique Application ID) is text
    for df in (summary, eligible, noneligible):
        first_col = df.columns[0]
        df[first_col] = df[first_col].astype(str).apply(lambda x: x.zfill(3) if x.isdigit() else x)
    
    try:
        summary.to_excel(summary_path, index=False)
        eligible.to_excel(eligible_path, index=False)
        noneligible.to_excel(noneligible_path, index=False)

        messagebox.showinfo(
            "Eligibility Summary Saved",
            f"✅ Files saved successfully:\n\n"
            f"📊 Total applications: {len(indata)}\n"
            f"✅ Eligible: {len(eligible)}\n"
            f"❌ Not eligible: {len(noneligible)}\n"
            f"📦 Total processed: {len(eligible) + len(noneligible)}"
        )

    except PermissionError:
        messagebox.showerror(
            "File Save Error",
            "⚠️ Cannot save one or more files.\n\n"
            "Please close them in Excel and try again."
        )
        return None


def generate_application_cards_and_marking(
    responses_path,
    output_folder,
    group_col,
    columns_to_select,
    num_apps=None,
    topics_col=None,
    num_scorers=3,
    unsure_options=None,
    no_options=None,
    progress_callback=None
):
    """
    Combined version of create_application_cards + eligibility marking sheet generator.
    Does NOT require the intermediate Application Summary.xlsx.
    """

    import pandas as pd
    import os
    from openpyxl import Workbook
    from openpyxl.worksheet.datavalidation import DataValidation
    from openpyxl.styles import PatternFill, Font, Alignment

    # === Step 1: Create application cards (Admin + Scorers) ===
    create_application_cards(
        responses_path=responses_path,
        columns_to_select=columns_to_select,
        output_folder=output_folder,
        group_col=group_col,
        progress_callback=progress_callback,
        num_apps=num_apps,
        topics_col=topics_col
    )

    # === Step 2: Build eligibility marking sheet directly ===
    df = pd.read_excel(responses_path)
    if num_apps is not None and isinstance(num_apps, int):
        df = df.head(num_apps)

    # Prepare data for marking sheet
    marking_headers = ["Unique Application ID", group_col]
    if topics_col:
        marking_headers.append(topics_col)

    for i in range(1, num_scorers + 1):
        marking_headers += [
            f"Scorer {i} Name",
            "Yes/ No/ Unsure (select from list)",
            "If Unsure, tell us the reason why (select from list)",
            "If No, tell us the reason why (select from list)",
            "Feedback"
        ]

    marking_data = []
    for idx, row in df.iterrows():
        uid = f"{idx+1:03d}"
        org = str(row[group_col])
        row_data = [uid, org]
        if topics_col:
            row_data.append(str(row.get(topics_col, "")))
        row_data += [""] * (len(marking_headers) - len(row_data))
        marking_data.append(row_data)

    # === Create workbook ===
    wb = Workbook()
    ws = wb.active
    ws.title = "Eligibility_Marking_Sheet"

    # Write headers
    for col_num, header in enumerate(marking_headers, start=1):
        cell = ws.cell(row=1, column=col_num, value=header)
        cell.font = Font(bold=True, name="Arial")
        cell.fill = PatternFill(start_color="FFFFCC", end_color="FFFFCC", fill_type="solid")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    # Write data
    for r, row in enumerate(marking_data, start=2):
        for c, val in enumerate(row, start=1):
            ws.cell(row=r, column=c, value=val)

    # === Dropdowns for Yes/No/Unsure ===
    yesno_dv = DataValidation(type="list", formula1='"Yes,No,Unsure"', allow_blank=True)
    ws.add_data_validation(yesno_dv)

    unsure_options = unsure_options or ["Lack of info online", "Some commercial aims", "Other"]
    no_options = no_options or ["Charity model", "Too new", "Other"]

    # Hidden list sheet
    hidden = wb.create_sheet("_lists")
    hidden.sheet_state = "hidden"
    for i, val in enumerate(unsure_options, start=1):
        hidden.cell(row=i, column=1, value=val)
    for i, val in enumerate(no_options, start=1):
        hidden.cell(row=i, column=2, value=val)

    unsure_dv = DataValidation(type="list", formula1=f"=_lists!$A$1:$A${len(unsure_options)}")
    no_dv = DataValidation(type="list", formula1=f"=_lists!$B$1:$B${len(no_options)}")

    ws.add_data_validation(unsure_dv)
    ws.add_data_validation(no_dv)

    # Attach dropdowns for each scorer section
    fixed_cols = 2 + (1 if topics_col else 0)
    for row in range(2, len(marking_data) + 2):
        for s in range(num_scorers):
            base = fixed_cols + s * 5
            yesno_dv.add(ws.cell(row=row, column=base + 2))
            unsure_dv.add(ws.cell(row=row, column=base + 3))
            no_dv.add(ws.cell(row=row, column=base + 4))

    # Save file
    marking_path = os.path.join(output_folder, "Eligibility Marking Sheet.xlsx")
    wb.save(marking_path)

    print(f"✅ Application cards and marking sheet created in {output_folder}")



